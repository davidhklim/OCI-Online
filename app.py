import os
import tempfile
import zipfile
from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename
from docx import Document
from docx2pdf import convert
import glob

from firms_data import FIRM_DATA

# Tell Flask that both templates and static files are in the current dir
app = Flask(
    __name__,
    static_folder='.',        # serve CSS/JS from project root
    static_url_path='/static',
    template_folder='.'       # look for index.html here
)
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()


def merge_fields(doc: Document, data: dict) -> Document:
    prefix, suffix = '«', '»'
    def replace_in_runs(runs):
        for run in runs:
            text = run.text
            for key, val in data.items():
                placeholder = f'{prefix}{key}{suffix}'
                if placeholder in text:
                    text = text.replace(placeholder, val)
            run.text = text
    for para in doc.paragraphs:
        replace_in_runs(para.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs)
    return doc


@app.route('/')
def home():
    return render_template('index.html')


@app.route('/firms', methods=['GET'])
def get_firms():
    def city_key(f): return f['City'].split(',')[0].strip()
    sorted_firms = sorted(FIRM_DATA, key=lambda f: (city_key(f), f['Firm']))
    return jsonify(sorted_firms)


@app.route('/upload-template', methods=['POST'])
def upload_template():
    if 'template' not in request.files:
        return "No template uploaded", 400
    f = request.files['template']
    filename = secure_filename(f.filename)
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    f.save(path)
    return jsonify({"template_path": path})


@app.route('/preview', methods=['POST'])
def preview_letters():
    data = request.get_json()
    tpl = data.get('template_path')
    sel = data.get('selected_firms', [])
    if not tpl or not os.path.exists(tpl) or not sel:
        return "Template not found or no firms selected", 400

    # Preview the first selected firm
    firm = sel[0]
    doc = Document(tpl)
    doc = merge_fields(doc, firm)

    base = os.path.splitext(os.path.basename(tpl))[0]
    tmp_docx = os.path.join(app.config['UPLOAD_FOLDER'], f"{base}_PREVIEW.docx")
    doc.save(tmp_docx)

    pdf_path = tmp_docx.replace('.docx', '.pdf')
    try:
        convert(tmp_docx, pdf_path)
    except Exception as e:
        print(f"Preview conversion error: {e}")

    if not os.path.exists(pdf_path):
        return "Preview PDF not generated", 500

    return jsonify({"url": f"/preview_pdf/{os.path.basename(pdf_path)}"})


@app.route('/preview_pdf/<filename>')
def serve_preview_pdf(filename):
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(path):
        return "Not found", 404
    return send_file(path, mimetype='application/pdf')


@app.route('/generate', methods=['POST'])
def generate_letters():
    data = request.get_json()
    tpl = data.get('template_path')
    sel = data.get('selected_firms', [])
    generate_pdf = data.get('generate_pdf', True)
    if not tpl or not os.path.exists(tpl):
        return "Template not found", 400

    def city_group(c):
        c0 = (c or '').split(',')[0].strip().lower()
        if c0 == 'toronto':   return 'Toronto'
        if c0 == 'vancouver': return 'Vancouver'
        if c0 == 'new york':  return 'New York'
        return 'Other'

    base, ext = os.path.splitext(os.path.basename(tpl))
    zip_path = tempfile.NamedTemporaryFile(delete=False, suffix='.zip').name

    # 1. Generate all DOCX files first
    city_dirs = set()
    for firm in sel:
        city = city_group(firm.get('City'))
        city_dir = os.path.join(app.config['UPLOAD_FOLDER'], city)
        city_dirs.add(city_dir)
        os.makedirs(city_dir, exist_ok=True)

        name_docx = f"{base} ({firm['Short_Name']}){ext}"
        path_docx = os.path.join(city_dir, name_docx)

        doc = Document(tpl)
        doc = merge_fields(doc, firm)
        doc.save(path_docx)

    # 2. Batch convert all DOCX to PDF per city
    if generate_pdf:
        for city_dir in city_dirs:
            docx_files = glob.glob(os.path.join(city_dir, '*.docx'))
            if docx_files:
                try:
                    convert(city_dir)  # Converts all DOCX in the directory to PDF
                except Exception as e:
                    print(f"Batch conversion error in {city_dir}: {e}")

    # 3. Add all DOCX and PDF files to the ZIP
    with zipfile.ZipFile(zip_path, 'w') as zf:
        for city_dir in city_dirs:
            city = os.path.basename(city_dir)
            for file in os.listdir(city_dir):
                if file.endswith('.docx') or (generate_pdf and file.endswith('.pdf')):
                    zf.write(os.path.join(city_dir, file), arcname=os.path.join(city, file))

    return send_file(zip_path, as_attachment=True, download_name='cover_letters.zip')


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
